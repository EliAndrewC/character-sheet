"""Character-sheet ingestion: file bytes in, plain text out.

Deterministic, no LLM involvement. Given a user-uploaded document this module:

  1. Enforces the upload size cap (design doc §5.4, default 1 MB).
  2. Detects the format from bytes (libmagic) - never trusts client-supplied
     content type or extension.
  3. Dispatches to a per-format extractor and returns an ``IngestResult``
     with the extracted plain text plus flags.

No LLM call happens here; that is Phase 4. The only job of this module is
"produce a clean text blob that a later stage can feed to Gemini."

See also ``import_url.py`` for URL fetching (which ultimately funnels back
into ``ingest_bytes`` once the URL is dereferenced).
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
from dataclasses import dataclass, field
from typing import Callable, Dict, List, Optional

import magic


# ---------------------------------------------------------------------------
# Config (env-overridable, design doc §5.4)
# ---------------------------------------------------------------------------

def _env_int(name: str, default: int) -> int:
    try:
        return int(os.environ.get(name, default))
    except ValueError:  # pragma: no cover - malformed env is a deploy bug
        return default


IMPORT_MAX_UPLOAD_MB = _env_int("IMPORT_MAX_UPLOAD_MB", 1)
IMPORT_MAX_UPLOAD_BYTES = IMPORT_MAX_UPLOAD_MB * 1024 * 1024

# A PDF whose text extraction yields fewer than this many characters per
# page is treated as "image-heavy" and flagged for multimodal fallback in
# Phase 4. Tuned small because character sheets have lots of labels.
NEAR_EMPTY_PDF_CHARS_PER_PAGE = 40

# PDFs larger than this many pages get capped when routed to multimodal
# (design §6.1). Phase 3 just records the cap; Phase 4 applies it.
IMPORT_MAX_PDF_PAGES = _env_int("IMPORT_MAX_PDF_PAGES", 10)


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------


class ImportIngestError(Exception):
    """Base class for every failure the ingest layer can raise."""

    error_code: str = "ingest_error"
    user_message: str = "We could not import this document."


class FileTooLargeError(ImportIngestError):
    error_code = "file_too_large"

    def __init__(self, size_bytes: int):
        self.size_bytes = size_bytes
        mb = size_bytes / (1024 * 1024)
        self.user_message = (
            f"This file is {mb:.1f} MB. The upload limit is "
            f"{IMPORT_MAX_UPLOAD_MB} MB, so this file is too large. "
            "Character sheets are small documents - if your file is "
            "much bigger than a megabyte, it probably contains images "
            "or other content that the importer does not use."
        )
        super().__init__(self.user_message)


class UnsupportedFormatError(ImportIngestError):
    error_code = "unsupported_format"

    def __init__(self, detected_mime: str):
        self.detected_mime = detected_mime
        self.user_message = (
            f"Unrecognised document format ({detected_mime}). Supported "
            "formats are: plain text, Markdown, RTF, HTML, PDF, "
            "Word (.doc/.docx), Excel (.xls/.xlsx), "
            "LibreOffice / OpenOffice (.odt/.ods/.sxw), and public "
            "Google Docs / Sheets URLs."
        )
        super().__init__(self.user_message)


class ParseError(ImportIngestError):
    error_code = "file_parse_error"

    def __init__(self, fmt: str, detail: str = ""):
        self.fmt = fmt
        self.detail = detail or "The file may be corrupted."
        self.user_message = f"Couldn't parse this file as {fmt}. {self.detail}"
        super().__init__(self.user_message)


class DocumentUnreadableError(ImportIngestError):
    """Raised when BOTH text extraction AND (later) multimodal fail.

    Phase 3 only raises this for ``.sxw`` files if ``odfpy`` cannot open
    them. Phase 4 reuses this error for the multimodal-fallback failure.
    """

    error_code = "document_unreadable"
    user_message = (
        "We could not read this document. Try exporting or re-saving it "
        "as plain text or .docx."
    )


# ---------------------------------------------------------------------------
# Result container
# ---------------------------------------------------------------------------


@dataclass
class IngestResult:
    text: str
    fmt: str                  # canonical format id ("txt", "pdf", ...)
    warnings: List[str] = field(default_factory=list)
    # True when a PDF returned near-empty text and Phase 4 should try the
    # multimodal-vision path instead of sending the text to the LLM.
    needs_multimodal_fallback: bool = False
    # Raw PDF bytes are kept around only when multimodal is needed; Phase 4
    # renders pages from these bytes.
    pdf_bytes_for_multimodal: Optional[bytes] = None


# ---------------------------------------------------------------------------
# Format detection
# ---------------------------------------------------------------------------

# Mapping from libmagic MIME to our canonical format id. Extensions are used
# only as a tiebreaker (e.g. plain text could be `.txt` or `.md`).
_MIME_TO_FMT: Dict[str, str] = {
    "text/plain": "txt",
    "text/markdown": "md",
    "text/x-markdown": "md",
    "text/html": "html",
    "text/xml": "html",  # HTML often detected as XML for well-formed pages
    "application/xhtml+xml": "html",
    "text/rtf": "rtf",
    "application/rtf": "rtf",
    "application/pdf": "pdf",
    "application/msword": "doc",
    "application/vnd.ms-excel": "xls",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        "xlsx",
    "application/vnd.oasis.opendocument.text": "odt",
    "application/vnd.oasis.opendocument.spreadsheet": "ods",
    "application/vnd.sun.xml.writer": "sxw",
    "application/csv": "csv",
    "text/csv": "csv",
}

# OOXML / ODF containers land as zip archives under libmagic. We peek at the
# bytes to split between docx, xlsx, odt, ods (all zip-based).
_ZIP_CONTAINER_MIMES = {"application/zip", "application/x-zip-compressed"}

# OLE compound documents (legacy .doc, .xls, .sxw all share magic bytes).
_OLE_COMPOUND_MIMES = {
    "application/x-ole-storage",
    "application/CDFV2",
    "application/vnd.ms-office",
}


def detect_format(data: bytes, filename: Optional[str] = None) -> str:
    """Return the canonical format id for ``data``.

    Uses libmagic as the primary signal and falls back to extension hints
    for cases libmagic reports generically (zip container, OLE compound).
    Raises ``UnsupportedFormatError`` if we don't handle the format.
    """
    mime = magic.from_buffer(data, mime=True) or ""
    ext = ""
    if filename:
        ext = os.path.splitext(filename)[1].lower().lstrip(".")

    if mime in _MIME_TO_FMT:
        fmt = _MIME_TO_FMT[mime]
        # text/plain could also be markdown or csv - the extension tiebreaker
        # only upgrades; it never downgrades a more-specific mime.
        if fmt == "txt" and ext in ("md", "csv"):
            return ext
        return fmt

    if mime in _ZIP_CONTAINER_MIMES or mime.startswith("application/zip"):
        fmt = _disambiguate_zip(data, ext)
        if fmt:
            return fmt

    if mime in _OLE_COMPOUND_MIMES:
        fmt = _disambiguate_ole(ext)
        if fmt:
            return fmt

    # Last resort: extension alone. Used when libmagic gives a generic
    # "application/octet-stream" for a file we can still handle.
    if ext in {"txt", "md", "html", "htm", "rtf", "pdf", "doc", "docx",
               "xls", "xlsx", "odt", "ods", "sxw", "csv"}:
        return "htm" if ext == "htm" else ext

    raise UnsupportedFormatError(mime or "unknown")


def _disambiguate_zip(data: bytes, ext: str) -> Optional[str]:
    """Peek into a zip container to tell OOXML from ODF."""
    try:
        import zipfile
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            names = set(zf.namelist())
            if "word/document.xml" in names:
                return "docx"
            if "xl/workbook.xml" in names:
                return "xlsx"
            if "mimetype" in names:
                mt = zf.read("mimetype").decode("ascii", errors="ignore")
                if "opendocument.text" in mt:
                    return "odt"
                if "opendocument.spreadsheet" in mt:
                    return "ods"
                if "sun.xml.writer" in mt:  # pragma: no cover - no .sxw fixture
                    return "sxw"
    except (zipfile.BadZipFile, KeyError):  # pragma: no cover - malformed zip
        return None
    # Fall back to the extension if we still can't tell.
    if ext in {"docx", "xlsx", "odt", "ods", "sxw"}:
        return ext
    return None  # pragma: no cover - every zip we see has markers OR an ext


def _disambiguate_ole(ext: str) -> Optional[str]:
    """Legacy .doc and .xls both look like OLE compound documents."""
    if ext in {"doc", "xls"}:
        return ext
    # Default to .doc since that's more common for L7R character sheets.
    return "doc"


# ---------------------------------------------------------------------------
# Per-format extractors
#
# Each extractor takes raw ``bytes`` and returns plain text. Failures
# produce ``ParseError`` (or ``DocumentUnreadableError`` when the format
# is known but the specific file just cannot be read).
# ---------------------------------------------------------------------------


def _extract_txt(data: bytes) -> str:
    # Try UTF-8, fall back to latin-1 which can't raise.
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1")


def _extract_md(data: bytes) -> str:
    # Markdown is readable as plain text; no conversion needed for LLM input.
    return _extract_txt(data)


def _extract_csv(data: bytes) -> str:
    return _extract_txt(data)


def _extract_rtf(data: bytes) -> str:
    from striprtf.striprtf import rtf_to_text
    try:
        return rtf_to_text(_extract_txt(data))
    except Exception as exc:
        raise ParseError("RTF", str(exc)) from exc


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup
    # Use lxml-fast when available but default html.parser to avoid adding
    # lxml as a hard dependency. We strip scripts/styles because embedded
    # JavaScript or CSS is a prompt-injection vector (design §5.1).
    soup = BeautifulSoup(_extract_txt(data), "html.parser")
    for tag in soup(["script", "style", "noscript", "template"]):
        tag.decompose()
    # Comments can also carry injected instructions; strip them.
    from bs4 import Comment
    for comment in soup.find_all(string=lambda t: isinstance(t, Comment)):
        comment.extract()
    return soup.get_text(separator="\n").strip()


def _extract_pdf(data: bytes) -> IngestResult:
    """PDF gets special treatment: we return an IngestResult directly so we
    can attach the near-empty flag and the raw bytes for multimodal."""
    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = reader.pages
        parts = []
        for page in pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        text = "\n".join(parts).strip()
        page_count = max(len(pages), 1)
    except Exception as exc:
        raise ParseError("PDF", str(exc)) from exc

    near_empty = len(text) / page_count < NEAR_EMPTY_PDF_CHARS_PER_PAGE
    warnings: List[str] = []
    if near_empty:
        warnings.append(
            "PDF text extraction produced very little content; Phase 4 "
            "will route this through the multimodal vision path."
        )
    return IngestResult(
        text=text,
        fmt="pdf",
        warnings=warnings,
        needs_multimodal_fallback=near_empty,
        pdf_bytes_for_multimodal=data if near_empty else None,
    )


def _extract_docx(data: bytes) -> str:
    from docx import Document
    try:
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise ParseError(".docx", str(exc)) from exc
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_doc(data: bytes) -> str:
    """Legacy .doc via antiword (design §6). Shells out to the binary."""
    antiword = shutil.which("antiword")
    if antiword is None:
        raise ParseError(
            ".doc",
            "antiword is not installed on this server; "
            "please re-save your document as .docx and try again.",
        )
    try:
        result = subprocess.run(
            [antiword, "-"],
            input=data,
            capture_output=True,
            timeout=30,
            check=True,
        )
    except subprocess.TimeoutExpired as exc:  # pragma: no cover - hardware-dep
        raise ParseError(".doc", "antiword timed out") from exc
    except subprocess.CalledProcessError as exc:
        raise ParseError(".doc", exc.stderr.decode("utf-8", "ignore")) from exc
    return result.stdout.decode("utf-8", errors="replace").strip()


def _extract_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook
    try:
        wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    except Exception as exc:
        raise ParseError(".xlsx", str(exc)) from exc
    parts: List[str] = []
    for ws in wb.worksheets:
        parts.append(f"### Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _extract_xls(data: bytes) -> str:
    import xlrd  # xlrd<2.0 required for .xls support
    try:
        wb = xlrd.open_workbook(file_contents=data)
    except Exception as exc:
        raise ParseError(".xls", str(exc)) from exc
    parts: List[str] = []
    for sheet in wb.sheets():
        parts.append(f"### Sheet: {sheet.name}")
        for rx in range(sheet.nrows):
            cells = [
                str(sheet.cell_value(rx, cx)) if sheet.cell_value(rx, cx) != "" else ""
                for cx in range(sheet.ncols)
            ]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _extract_odt(data: bytes) -> str:
    return _extract_odf_text(data, ".odt")


def _extract_sxw(data: bytes) -> str:
    # odfpy's public API handles .sxw the same as .odt for reading purposes.
    return _extract_odf_text(data, ".sxw")


def _extract_ods(data: bytes) -> str:
    from odf.opendocument import load
    from odf.table import Table, TableRow, TableCell
    from odf import teletype
    try:
        doc = load(io.BytesIO(data))
    except Exception as exc:
        raise ParseError(".ods", str(exc)) from exc
    parts: List[str] = []
    for table in doc.spreadsheet.getElementsByType(Table):
        parts.append(f"### Sheet: {table.getAttribute('name')}")
        for row in table.getElementsByType(TableRow):
            cells = [
                teletype.extractText(cell)
                for cell in row.getElementsByType(TableCell)
            ]
            if any(c.strip() for c in cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _extract_odf_text(data: bytes, label: str) -> str:
    from odf.opendocument import load
    from odf.text import P
    from odf import teletype
    try:
        doc = load(io.BytesIO(data))
    except Exception as exc:
        raise DocumentUnreadableError(
            f"Could not open {label} document: {exc}"
        ) from exc
    parts: List[str] = []
    for p in doc.getElementsByType(P):
        parts.append(teletype.extractText(p))
    return "\n".join(parts).strip()


_EXTRACTORS: Dict[str, Callable[[bytes], "IngestResult | str"]] = {
    "txt": _extract_txt,
    "md": _extract_md,
    "csv": _extract_csv,
    "rtf": _extract_rtf,
    "html": _extract_html,
    "pdf": _extract_pdf,
    "doc": _extract_doc,
    "docx": _extract_docx,
    "xls": _extract_xls,
    "xlsx": _extract_xlsx,
    "odt": _extract_odt,
    "ods": _extract_ods,
    "sxw": _extract_sxw,
}


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def ingest_bytes(
    data: bytes,
    filename: Optional[str] = None,
    *,
    enforce_size_limit: bool = True,
) -> IngestResult:
    """Take raw file bytes, return extracted plain text + metadata.

    Raises subclasses of ``ImportIngestError`` on failure. Size-cap
    enforcement can be disabled (e.g. when the URL fetcher has already
    enforced it) but by default is on.
    """
    if enforce_size_limit and len(data) > IMPORT_MAX_UPLOAD_BYTES:
        raise FileTooLargeError(len(data))

    fmt = detect_format(data, filename=filename)
    extractor = _EXTRACTORS.get(fmt)
    if extractor is None:  # pragma: no cover - unreachable: detect_format raises
        raise UnsupportedFormatError(fmt)

    result = extractor(data)
    if isinstance(result, IngestResult):
        return result
    return IngestResult(text=result, fmt=fmt)


__all__ = [
    "IngestResult",
    "ImportIngestError",
    "FileTooLargeError",
    "UnsupportedFormatError",
    "ParseError",
    "DocumentUnreadableError",
    "ingest_bytes",
    "detect_format",
    "IMPORT_MAX_UPLOAD_BYTES",
    "IMPORT_MAX_UPLOAD_MB",
    "NEAR_EMPTY_PDF_CHARS_PER_PAGE",
    "IMPORT_MAX_PDF_PAGES",
]
